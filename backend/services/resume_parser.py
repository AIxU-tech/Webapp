"""
Resume Parser Service

Uses Claude AI (Haiku 4.5) to extract structured profile information from resumes.
Supports PDF and Word documents uploaded to GCS.

The parsing runs in a background thread after being triggered by the user.
On completion (or failure), a WebSocket event is emitted to the user.
"""

import base64
import json
import threading
from datetime import date

import anthropic

from backend.extensions import db
from backend.models.user import User
from backend.models.profile_sections import Education, Experience, Project
from backend.services.storage import _get_bucket
from backend.sockets.events import emit_resume_parse_result


# Model for resume parsing
PARSE_MODEL = "claude-haiku-4-5-20251001"
PARSE_MAX_TOKENS = 4096

# Tracks user_ids whose resume is currently being parsed. Prevents a user from
# kicking off multiple concurrent parses (which would waste Claude tokens and
# race on the same User row). This is per-process in-memory state; it matches
# the semantics of the previous polling-based implementation.
_active_parses = set()
_active_parses_lock = threading.Lock()


def try_acquire_parse_slot(user_id):
    """
    Atomically reserve a parse slot for ``user_id``.

    Returns True if the slot was acquired (caller is responsible for releasing
    it, either directly on failure or indirectly via the background thread on
    success). Returns False if a parse is already in progress for this user.
    """
    with _active_parses_lock:
        if user_id in _active_parses:
            return False
        _active_parses.add(user_id)
        return True


def release_parse_slot(user_id):
    """Release a parse slot. Safe to call even if the slot is not held."""
    with _active_parses_lock:
        _active_parses.discard(user_id)


def download_file_from_gcs(gcs_path):
    """Download a file from GCS and return its bytes."""
    bucket = _get_bucket()
    blob = bucket.blob(gcs_path)
    return blob.download_as_bytes()


def _truncate(value, max_len):
    """Truncate a string to fit a DB column, or return None for non-strings."""
    if not value or not isinstance(value, str):
        return value
    return value[:max_len]


def _extract_text_from_docx(file_bytes):
    """Extract text from a Word document (.docx), including table content."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


SYSTEM_PROMPT = """You are a resume parser that extracts structured profile information. Your output will be fed directly into a database, so quality and precision matter.

Think carefully before outputting each field. Each field has a specific purpose — do NOT duplicate information across fields.

Return a JSON object with this structure:

{
    "about_section": "A concise 2-3 sentence professional summary synthesized from the resume. Do NOT just copy-paste bullet points or list awards.",
    "headline": "Short professional headline (e.g. 'Software Engineer | AI Researcher'). Max ~10 words.",
    "location": "City, State or City, Country if found, or null",
    "skills": ["skill1", "skill2", ...],
    "social_links": [{"type": "linkedin", "url": "..."}, ...],
    "education": [
        {
            "institution": "University Name",
            "degree": "Degree type only (e.g. 'Bachelor of Science', 'Master of Arts')",
            "field_of_study": "Major/field only (e.g. 'Computer Science'). Do NOT repeat the degree type here.",
            "start_date": "YYYY-MM-DD or null",
            "end_date": "YYYY-MM-DD or null",
            "gpa": "Numeric GPA value only (e.g. '3.8') or null. Do not include the scale.",
            "description": "Brief notable achievements ONLY if highly relevant (e.g. 'Magna Cum Laude'). Do NOT dump a raw list of awards, honors, or coursework. Null if nothing notable."
        }
    ],
    "experience": [
        {
            "title": "Job Title",
            "company": "Company Name",
            "location": "City, State or null",
            "start_date": "YYYY-MM-DD",
            "end_date": "YYYY-MM-DD or null",
            "description": "Concise summary of key responsibilities and achievements. Write in clean prose or short bullet points, not a raw copy-paste from the resume."
        }
    ],
    "projects": [
        {
            "title": "Project Name",
            "description": "What the project does and your role. Keep it concise.",
            "url": "Project URL or null",
            "technologies": ["tech1", "tech2"]
        }
    ]
}

Critical rules:
- NO DUPLICATION: Each piece of information goes in exactly one field. If a major goes in field_of_study, do NOT also put it in degree or description.
- CLEAN OUTPUT: Description fields should read naturally, not look like raw resume dumps. Rewrite if needed.
- EDUCATION DESCRIPTIONS: Only include genuinely notable honors (Dean's List, Latin honors, Phi Beta Kappa). Do NOT list every award, scholarship, or coursework item. If there's nothing standout, use null.
- DEGREE vs FIELD_OF_STUDY: "degree" is the credential type (Bachelor of Science, MBA). "field_of_study" is the subject (Computer Science, Finance). Keep them separate.
- For dates: YYYY-MM-DD format. Year only → YYYY-01-01. Month+year → YYYY-MM-01. Current/present → null for end_date.
- For social_links: supported types are linkedin, x, instagram, github, discord, youtube, website
- Return ONLY the JSON object, no other text."""


def _parse_with_claude(file_bytes, content_type):
    """Send resume content to Claude and get structured data back."""
    client = anthropic.Anthropic()

    if content_type == "application/pdf":
        content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": base64.b64encode(file_bytes).decode("utf-8"),
                },
            },
            {
                "type": "text",
                "text": "Parse this resume and extract all structured information. Return only JSON.",
            },
        ]
    elif content_type == "application/msword":
        raise ValueError(
            "Old .doc format is not supported. Please convert to .docx or PDF."
        )
    else:
        text = _extract_text_from_docx(file_bytes)
        content = (
            "Parse the resume text between <resume> tags and extract all "
            "structured information. Return only JSON.\n\n"
            f"<resume>\n{text}\n</resume>"
        )

    response = client.messages.create(
        model=PARSE_MODEL,
        max_tokens=PARSE_MAX_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": content}],
        timeout=60.0,
    )

    from backend.services.agents.base import extract_json

    parsed = extract_json(response.content)
    if not isinstance(parsed, dict):
        raise ValueError("Expected JSON object from resume parse, got array")
    return parsed


def _parse_date_safe(value):
    """Parse a date string safely, returning None on failure."""
    if not value:
        return None
    try:
        return date.fromisoformat(value)
    except (ValueError, TypeError):
        return None


def _apply_parsed_data(app, user_id, parsed_data):
    """Apply parsed resume data to user profile within app context."""
    with app.app_context():
        try:
            user = db.session.get(User, user_id)
            if not user:
                return

            # Update text profile fields only if currently empty
            if not user.about_section and parsed_data.get("about_section"):
                user.about_section = parsed_data["about_section"]

            if not user.headline and parsed_data.get("headline"):
                user.headline = _truncate(parsed_data["headline"], 200)

            if not user.location and parsed_data.get("location"):
                user.location = _truncate(parsed_data["location"], 100)

            # Merge skills (union of existing + new)
            if parsed_data.get("skills"):
                existing_skills = []
                if user.skills:
                    try:
                        existing_skills = json.loads(user.skills)
                    except (json.JSONDecodeError, TypeError):
                        existing_skills = []
                merged = list(dict.fromkeys(existing_skills + parsed_data["skills"]))
                user.skills = json.dumps(merged)

            # Merge social links (add types not already present)
            if parsed_data.get("social_links"):
                existing_links = user.get_social_links_list()
                existing_types = {link["type"] for link in existing_links}
                for link in parsed_data["social_links"]:
                    if (
                        link.get("type")
                        and link.get("url")
                        and link["type"] not in existing_types
                    ):
                        existing_links.append(link)
                user.set_social_links_list(existing_links)

            # Add education entries (skip duplicates by institution+degree)
            for edu_data in parsed_data.get("education", []):
                if not edu_data.get("institution") or not edu_data.get("degree"):
                    continue
                exists = Education.query.filter_by(
                    user_id=user_id,
                    institution=edu_data["institution"],
                    degree=edu_data["degree"],
                ).first()
                if exists:
                    continue
                gpa_raw = edu_data.get("gpa")
                db.session.add(
                    Education(
                        user_id=user_id,
                        institution=_truncate(edu_data["institution"], 200),
                        degree=_truncate(edu_data["degree"], 200),
                        field_of_study=_truncate(edu_data.get("field_of_study"), 200),
                        start_date=_parse_date_safe(edu_data.get("start_date")),
                        end_date=_parse_date_safe(edu_data.get("end_date")),
                        gpa=str(gpa_raw)[:20] if gpa_raw is not None else None,
                        description=edu_data.get("description"),
                    )
                )

            # Add experience entries (skip duplicates by title+company)
            for exp_data in parsed_data.get("experience", []):
                if not exp_data.get("title") or not exp_data.get("company"):
                    continue
                start = _parse_date_safe(exp_data.get("start_date"))
                if not start:
                    continue  # start_date is required
                exists = Experience.query.filter_by(
                    user_id=user_id,
                    title=exp_data["title"],
                    company=exp_data["company"],
                ).first()
                if exists:
                    continue
                db.session.add(
                    Experience(
                        user_id=user_id,
                        title=_truncate(exp_data["title"], 200),
                        company=_truncate(exp_data["company"], 200),
                        location=_truncate(exp_data.get("location"), 200),
                        start_date=start,
                        end_date=_parse_date_safe(exp_data.get("end_date")),
                        description=exp_data.get("description"),
                    )
                )

            # Add project entries (skip duplicates by title)
            for proj_data in parsed_data.get("projects", []):
                if not proj_data.get("title"):
                    continue
                exists = Project.query.filter_by(
                    user_id=user_id,
                    title=proj_data["title"],
                ).first()
                if exists:
                    continue
                techs = proj_data.get("technologies")
                db.session.add(
                    Project(
                        user_id=user_id,
                        title=_truncate(proj_data["title"], 200),
                        description=proj_data.get("description"),
                        url=_truncate(proj_data.get("url"), 500),
                        technologies=json.dumps(techs) if isinstance(techs, list) else None,
                    )
                )

            db.session.commit()
        except Exception:
            db.session.rollback()
            raise


def start_resume_parse(app, user_id, file_bytes, content_type):
    """
    Start resume parsing in a background thread.

    The file bytes are downloaded in the request context before calling this.
    The thread uses app context only for database writes and the socket emit.
    """

    def _run():
        try:
            parsed_data = _parse_with_claude(file_bytes, content_type)
            _apply_parsed_data(app, user_id, parsed_data)
            with app.app_context():
                emit_resume_parse_result(user_id, success=True)
        except Exception as e:
            print(f"[ResumeParser] Error parsing resume for user {user_id}: {e}")
            with app.app_context():
                emit_resume_parse_result(
                    user_id,
                    success=False,
                    error_message="Failed to parse resume. Please try again later.",
                )
        finally:
            release_parse_slot(user_id)

    thread = threading.Thread(target=_run, daemon=True)
    thread.start()
